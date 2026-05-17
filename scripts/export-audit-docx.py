#!/usr/bin/env python3
"""Convert the deep feature audit Markdown report to a .docx using python-docx.

Lightweight fallback when pandoc is not available. Headings (#, ##, ###),
paragraphs, bullet lists, numbered lists, and pipe tables are preserved.
Inline emphasis (**bold**, *italic*, `code`) is preserved at a basic level.

Usage:
    python3 scripts/export-audit-docx.py [INPUT.md] [OUTPUT.docx]

Defaults:
    INPUT   = docs/audits/XUANTOI_DEEP_FEATURE_AUDIT_2026_05_17.md
    OUTPUT  = same path with .docx extension
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    sys.stderr.write(
        "python-docx is not installed. Run: pip install python-docx\n"
    )
    sys.exit(2)


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
NUMBERED_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
INLINE_BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
INLINE_ITALIC_RE = re.compile(r"(?<![*\w])\*([^*]+)\*(?![*\w])")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")


def add_inline_runs(paragraph, text: str) -> None:
    """Apply minimal bold/italic/code formatting to a paragraph."""
    parts = [(text, set())]

    def split(pattern, parts, mark):
        out = []
        for chunk, marks in parts:
            if mark in marks:
                out.append((chunk, marks))
                continue
            last = 0
            for m in pattern.finditer(chunk):
                if m.start() > last:
                    out.append((chunk[last : m.start()], marks))
                out.append((m.group(1), marks | {mark}))
                last = m.end()
            if last < len(chunk):
                out.append((chunk[last:], marks))
        return out

    parts = split(INLINE_CODE_RE, parts, "code")
    parts = split(INLINE_BOLD_RE, parts, "bold")
    parts = split(INLINE_ITALIC_RE, parts, "italic")

    for chunk, marks in parts:
        run = paragraph.add_run(chunk)
        if "bold" in marks:
            run.bold = True
        if "italic" in marks:
            run.italic = True
        if "code" in marks:
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def flush_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    # rows[1] is the alignment row in GFM tables: |---|---|.
    if len(rows) >= 2 and all(re.fullmatch(r":?-+:?", c.strip()) for c in rows[1]):
        header = rows[0]
        body = rows[2:]
    else:
        header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Light List Accent 1"
    for col, cell in enumerate(header):
        table.cell(0, col).text = cell.strip()
    for r_idx, row in enumerate(body, start=1):
        for col in range(len(header)):
            cell_value = row[col].strip() if col < len(row) else ""
            table.cell(r_idx, col).text = cell_value


def parse_table_row(line: str) -> list[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [c.strip() for c in raw.split("|")]


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()

    pending_table: list[list[str]] = []

    for raw_line in md_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            if pending_table:
                flush_table(doc, pending_table)
                pending_table = []
            continue

        if TABLE_ROW_RE.match(line):
            pending_table.append(parse_table_row(line))
            continue

        if pending_table:
            flush_table(doc, pending_table)
            pending_table = []

        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            doc.add_heading(text, level=min(level, 4))
            continue

        m = BULLET_RE.match(line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, m.group(1))
            continue

        m = NUMBERED_RE.match(line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m.group(1))
            continue

        if line.startswith("```"):
            # Skip code fences entirely (best-effort fallback).
            continue
        if line.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline_runs(p, line.lstrip("> ").strip())
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, line)

    if pending_table:
        flush_table(doc, pending_table)

    doc.save(docx_path)


def main() -> None:
    args = sys.argv[1:]
    if len(args) >= 1:
        md_path = Path(args[0])
    else:
        md_path = Path("docs/audits/XUANTOI_DEEP_FEATURE_AUDIT_2026_05_17.md")
    if len(args) >= 2:
        docx_path = Path(args[1])
    else:
        docx_path = md_path.with_suffix(".docx")
    if not md_path.exists():
        sys.stderr.write(f"ERROR: input not found: {md_path}\n")
        sys.exit(1)
    convert(md_path, docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    main()
